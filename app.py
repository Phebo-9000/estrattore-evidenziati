import streamlit as st

# Password per proteggere l'accesso
PASSWORD = "Synergia.ai"

def check_password():
    """Controlla se la password inserita è corretta"""
    st.sidebar.header("🔒 Accesso Riservato")
    password = st.sidebar.text_input("Inserisci la password:", type="password")
    if password == PASSWORD:
        return True
    else:
        st.warning("❌ Password errata. Riprova.")
        return False

# Se la password è errata, non mostrare l'app
if not check_password():
    st.stop()
from docx import Document
import os
from collections import defaultdict
from datetime import datetime

def extract_highlighted_texts(doc, filename):
    """
    Estrae i testi evidenziati da un documento Word e li organizza per colore,
    mantenendo il riferimento al file originale.
    """
    highlighted_texts = defaultdict(lambda: defaultdict(list))
    
    for para in doc.paragraphs:
        temp_text = ""
        last_color = None
        
        for run in para.runs:
            if run.font.highlight_color:
                color = run.font.highlight_color
                if last_color is not None and color != last_color:
                    highlighted_texts[color][filename].append(temp_text.strip())
                    temp_text = ""
                temp_text += run.text.strip() + " "
                last_color = color
            elif last_color is not None:
                highlighted_texts[last_color][filename].append(temp_text.strip())
                temp_text = ""
                last_color = None
        
        if temp_text and last_color is not None:
            highlighted_texts[last_color][filename].append(temp_text.strip())
    
    return highlighted_texts

def extract_categories(doc):
    """
    Estrae i titoli e i loro colori di evidenziazione dal file delle categorie.
    """
    categories = {}
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color and run.text.strip():
                categories[run.font.highlight_color] = run.text.strip()
    
    return categories

def generate_report(highlighted_texts, categories, output_path):
    """
    Genera un documento Word con i testi organizzati per categoria,
    raggruppando gli estratti per file di origine.
    """
    doc = Document()
    
    for color, category in categories.items():
        doc.add_heading(category, level=1)
        
        if color in highlighted_texts:
            for filename, texts in highlighted_texts[color].items():
                doc.add_paragraph(f"**{filename}**", style='Heading 3')
                for text in texts:
                    doc.add_paragraph(text)
        else:
            doc.add_paragraph("⚠ Nessun testo corrispondente trovato per questa categoria.")
    
    unmatched_colors = set(highlighted_texts.keys()) - set(categories.keys())
    if unmatched_colors:
        doc.add_heading("⚠ Colori non corrispondenti trovati", level=1)
        for color in unmatched_colors:
            for filename, texts in highlighted_texts[color].items():
                doc.add_paragraph(f"**{filename}**", style='Heading 3')
                for text in texts:
                    doc.add_paragraph(text)
    
    doc.save(output_path)
    return output_path

# Interfaccia Web con Streamlit
st.title("📝 Estrattore di Testi Evidenziati")
st.write("Carica i documenti Word con testi evidenziati e il file delle categorie per generare un report.")

uploaded_files = st.file_uploader("Carica i file Word con testi evidenziati", accept_multiple_files=True, type=["docx"])
label_file = st.file_uploader("Carica il file delle categorie", type=["docx"])

if uploaded_files and label_file:
    highlighted_texts = defaultdict(lambda: defaultdict(list))
    
    for file in uploaded_files:
        doc = Document(file)
        extracted_texts = extract_highlighted_texts(doc, file.name)
        for color, files in extracted_texts.items():
            for filename, texts in files.items():
                highlighted_texts[color][filename].extend(texts)
    
    category_doc = Document(label_file)
    categories = extract_categories(category_doc)
    
    output_filename = f"Report_Evidenziati_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    generate_report(highlighted_texts, categories, output_filename)

    st.success("✅ Report generato con successo!")
    with open(output_filename, "rb") as f:
        st.download_button("📥 Scarica il Report", f, file_name=output_filename)
